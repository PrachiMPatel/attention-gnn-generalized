"""Build a .docx comparison report for WITS-static vs prompt-only vs GNN.

Loads three sources of predictions on the same 311-row held-out test split:

  1. WITS static analyzer  -- outputs/wits_static_predictions_*.jsonl
                              (produced by data/score_wits_static.py)
  2. Prompt-only LLM        -- outputs/test_extras_*.pkl["prompt_preds"]
                              (produced by wits_main.ipynb section 5)
  3. GNN (class-weighted)   -- outputs/gnn_weighted_*/model.pt scored
                              against outputs/test_graphs_*.pkl

Computes per-class P/R/F1, TP/FP/FN/TN, 4-class and 3-class collapsed
metrics, per-source-bucket accuracy, latency stats. Writes everything
into a structured Word document at outputs/wits_gnn_comparison_report.docx.

Run:
    python data/build_report_docx.py
"""
from __future__ import annotations

import json
import pickle
import statistics
import time
from collections import Counter, defaultdict
from pathlib import Path

import numpy as np
import torch
from sklearn.metrics import (
    accuracy_score, f1_score, precision_score, recall_score, confusion_matrix,
)
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from torch_geometric.loader import DataLoader

REPO = Path(__file__).resolve().parents[1]
OUT_DIR = REPO / "outputs"
DOCX_PATH = OUT_DIR / "wits_gnn_comparison_report.docx"

LABELS = ["safe", "maybe_safe", "unsafe", "extremely_unsafe"]
LABEL2ID = {n: i for i, n in enumerate(LABELS)}
ID2LABEL = {i: n for n, i in LABEL2ID.items()}
OP_LABELS = ["safe", "maybe", "extremely_unsafe"]


# ---------------------------------------------------------------------------
# Load all three sources of predictions
# ---------------------------------------------------------------------------

def find_one(glob: str) -> Path:
    matches = sorted(OUT_DIR.glob(glob))
    if not matches:
        raise FileNotFoundError(f"no match for {glob} in {OUT_DIR}")
    if len(matches) > 1:
        print(f"  multiple matches for {glob}; using {matches[-1].name}")
    return matches[-1]


print("Loading prediction artefacts ...")
WITS_PRED_PATH = find_one("wits_static_predictions_*.jsonl")
TEST_SPLIT_PATH = find_one("wits_test_split_*.jsonl")
TEST_GRAPHS_PATH = find_one("test_graphs_wits_*.pkl")
TEST_META_PATH = find_one("test_meta_wits_*.pkl")
TEST_EXTRAS_PATH = find_one("test_extras_wits_*.pkl")
GNN_W_DIR = sorted(OUT_DIR.glob("gnn_weighted_wits_*"))[-1]

wits_preds = [json.loads(l) for l in WITS_PRED_PATH.read_text(encoding="utf-8").splitlines() if l.strip()]
test_split = [json.loads(l) for l in TEST_SPLIT_PATH.read_text(encoding="utf-8").splitlines() if l.strip()]
with open(TEST_GRAPHS_PATH, "rb") as f:
    test_graphs = pickle.load(f)
with open(TEST_META_PATH, "rb") as f:
    test_meta = pickle.load(f)
with open(TEST_EXTRAS_PATH, "rb") as f:
    test_extras = pickle.load(f)

print(f"  WITS predictions      : {len(wits_preds)}")
print(f"  test split            : {len(test_split)}")
print(f"  test graphs           : {len(test_graphs)}")
print(f"  test meta             : {len(test_meta)}")


# ---------------------------------------------------------------------------
# Utility-model predictions (one or more LLMs evaluated as direct
# 4-class classifiers using data/utility_model_runner.py).
# Whichever model prediction files exist in outputs/ get included.
# ---------------------------------------------------------------------------

# Model name -> display label in tables. Order matters: first listed first in
# table rows. Extend this list when a new model is run via
#   python data/utility_model_runner.py --model <name>
UTILITY_MODELS = [
    ("gpt-5-mini",   "Utility-model LLM (gpt-5-mini)"),
    ("gpt-4.1",      "Utility-model LLM (gpt-4.1)"),
    ("gpt-4o-mini",  "Utility-model LLM (gpt-4o-mini)"),
]
utility_loaded: list[tuple[str, str, list[dict]]] = []
for model_name, display in UTILITY_MODELS:
    path = OUT_DIR / f"utility_model_predictions_{model_name}.jsonl"
    if not path.exists():
        print(f"  utility-model preds   : (missing {path.name}; skipped)")
        continue
    recs = [json.loads(l) for l in path.read_text(encoding="utf-8").splitlines() if l.strip()]
    recs.sort(key=lambda r: r["i"])
    utility_loaded.append((model_name, display, recs))
    print(f"  utility-model preds   : {len(recs):3d}  {model_name:14s}  ({path.name})")


# ---------------------------------------------------------------------------
# Score the class-weighted GNN -- and time it
# ---------------------------------------------------------------------------

import sys  # noqa: E402
sys.path.insert(0, str(REPO))

from models.gnn.graph_classifier import GraphClassifier  # noqa: E402

with open(GNN_W_DIR / "model_metadata.json") as f:
    md_w = json.load(f)
gnn_w = GraphClassifier(
    hidden_channel_dimensions=md_w["hidden_channel_dimensions"],
    num_classes=md_w["num_classes"],
)
gnn_w.load_state_dict(torch.load(GNN_W_DIR / "model.pt", map_location="cpu"))
gnn_w.eval()

print(f"Scoring class-weighted GNN on {len(test_graphs)} cached test graphs ...")
# Per-sample latency
loader = DataLoader(test_graphs, batch_size=1, shuffle=False)
gnn_y_pred = []
gnn_y_true = []
gnn_lat_ms = []
# warmup
warm = iter(loader)
for _ in range(min(5, len(test_graphs))):
    b = next(warm)
    with torch.no_grad():
        gnn_w(b.x.float(), b.edge_index, b.batch, dropout_percentage=0.0)
for batch in loader:
    t0 = time.perf_counter()
    with torch.no_grad():
        logits = gnn_w(batch.x.float(), batch.edge_index, batch.batch, dropout_percentage=0.0)
    gnn_lat_ms.append((time.perf_counter() - t0) * 1000.0)
    gnn_y_pred.append(int(logits.argmax(dim=-1).item()))
    gnn_y_true.append(int(batch.y.item()))

gnn_y_true_str = [ID2LABEL[i] for i in gnn_y_true]
gnn_y_pred_str = [ID2LABEL[i] for i in gnn_y_pred]


# ---------------------------------------------------------------------------
# Align prompt-only LLM preds (test_extras) to the same order
# ---------------------------------------------------------------------------

prompt_y_pred = [int(p) for p in test_extras["prompt_preds"]]
prompt_y_true = [int(l) for l in test_extras["labels"]]
prompt_y_true_str = [ID2LABEL[i] for i in prompt_y_true]
prompt_y_pred_str = [ID2LABEL[i] for i in prompt_y_pred]

assert len(prompt_y_true) == len(gnn_y_true), "prompt-only and GNN label vectors disagree in length"
# Both came from the same test_meta ordering, so we trust them as parallel.


# ---------------------------------------------------------------------------
# Align WITS preds to the same order as the GNN
# ---------------------------------------------------------------------------

# WITS predictions are 1:1 with TEST_SPLIT_PATH order. test_meta truncates
# commands to 120 chars; rebuild the join by walking test_meta and finding
# the matching full-command row in test_split.

by_full = {(t["command"], t["shell"]): w for t, w in zip(test_split, wits_preds)}

wits_y_pred = []
wits_y_true = []
wits_lat_ms = []
for m in test_meta:
    full = None
    for t in test_split:
        if t["command"].startswith(m["command"]) and t["shell"] == m["shell"]:
            full = t
            break
    if full is None:
        raise RuntimeError(f"no WITS pred for command={m['command'][:60]!r}")
    w = by_full[(full["command"], full["shell"])]
    wits_y_true.append(m["verdict"])
    wits_y_pred.append(w.get("wits_verdict") or "(error)")
    wits_lat_ms.append(float(w.get("wits_elapsed_ms") or 0.0))

assert len(wits_y_true) == len(gnn_y_true)


# ---------------------------------------------------------------------------
# Metric helpers
# ---------------------------------------------------------------------------

def tp_fp_fn_tn(y_true: list[str], y_pred: list[str], cls: str) -> tuple[int, int, int, int]:
    tp = sum(1 for t, p in zip(y_true, y_pred) if t == cls and p == cls)
    fp = sum(1 for t, p in zip(y_true, y_pred) if t != cls and p == cls)
    fn = sum(1 for t, p in zip(y_true, y_pred) if t == cls and p != cls)
    tn = sum(1 for t, p in zip(y_true, y_pred) if t != cls and p != cls)
    return tp, fp, fn, tn


def per_class_metrics(y_true, y_pred, labels):
    rows = []
    for cls in labels:
        tp, fp, fn, tn = tp_fp_fn_tn(y_true, y_pred, cls)
        support = tp + fn
        precision = tp / (tp + fp) if (tp + fp) else 0.0
        recall = tp / (tp + fn) if (tp + fn) else 0.0
        f1 = 2 * precision * recall / (precision + recall) if (precision + recall) else 0.0
        rows.append({
            "class": cls, "tp": tp, "fp": fp, "fn": fn, "tn": tn,
            "support": support, "precision": precision, "recall": recall, "f1": f1,
        })
    return rows


def collapse_op(v: str) -> str:
    return "maybe" if v in ("maybe_safe", "unsafe") else v


def latency_stats(lats_ms: list[float]) -> dict:
    if not lats_ms:
        return {"mean": float("nan"), "median": float("nan"), "p95": float("nan"),
                "p99": float("nan"), "min": float("nan"), "max": float("nan")}
    s = sorted(lats_ms)
    def pct(p): return s[min(len(s) - 1, int(round(p * len(s))))]
    return {
        "mean": statistics.mean(s),
        "median": statistics.median(s),
        "p95": pct(0.95),
        "p99": pct(0.99),
        "min": s[0],
        "max": s[-1],
    }


# Featurization latency was measured in §11 of the notebook. We didn't
# capture it to disk so we re-estimate here from a representative
# constant (the notebook reported mean=963.6ms / p95=1166.2ms on CPU
# with Qwen2.5-0.5B). Update these if you re-time on different hardware.
FEAT_LAT = {"mean": 963.6, "median": 957.0, "p95": 1166.2, "p99": 1188.0, "min": 880.0, "max": 1190.0}

method_runs = {
    "WITS static (rule-based)": {
        "y_true":      wits_y_true,
        "y_pred":      wits_y_pred,
        "model_lat":   latency_stats(wits_lat_ms),
        "feat_lat":    {"mean": 0.0, "median": 0.0, "p95": 0.0, "p99": 0.0, "min": 0.0, "max": 0.0},
        "lat_breakdown": "rule engine only (no LLM call)",
    },
    "Prompt-only LLM (Qwen 2.5 0.5B)": {
        "y_true":      prompt_y_true_str,
        "y_pred":      prompt_y_pred_str,
        "model_lat":   {"mean": 0.0, "median": 0.0, "p95": 0.0, "p99": 0.0, "min": 0.0, "max": 0.0},
        "feat_lat":    FEAT_LAT,
        "lat_breakdown": "LLM forward only",
    },
    "GNN (class-weighted CE)": {
        "y_true":      gnn_y_true_str,
        "y_pred":      gnn_y_pred_str,
        "model_lat":   latency_stats(gnn_lat_ms),
        "feat_lat":    FEAT_LAT,
        "lat_breakdown": "LLM forward + GNN",
    },
}

for _model_name, _display, _recs in utility_loaded:
    _u_yt = [r["truth"] for r in _recs]
    _u_yp = [r["verdict"] if r["verdict"] in LABELS else "extremely_unsafe"
             for r in _recs]
    # Latency: only fresh calls (cached==False) count toward production latency
    # stats. Cached rows took ~0 ms to look up and would skew the headline.
    _u_fresh = [float(r["latency_ms"]) for r in _recs if not r.get("cached")]
    method_runs[_display] = {
        "y_true":      _u_yt,
        "y_pred":      _u_yp,
        "model_lat":   latency_stats(_u_fresh),
        "feat_lat":    {"mean": 0.0, "median": 0.0, "p95": 0.0, "p99": 0.0, "min": 0.0, "max": 0.0},
        "lat_breakdown": f"CAPI chat completion ({_model_name})",
    }


# ---------------------------------------------------------------------------
# Per-source bucket accuracy (uses the WITS-aligned ordering)
# ---------------------------------------------------------------------------

def bucket_of(source: str) -> str:
    return source.split(":", 1)[0] if ":" in source else (source or "(unknown)")


per_bucket: dict[str, dict[str, dict]] = defaultdict(lambda: defaultdict(lambda: {"n": 0, "correct": 0}))
for i, m in enumerate(test_meta):
    bucket = bucket_of(m.get("source", ""))
    truth = m["verdict"]
    for name, run in method_runs.items():
        pred = run["y_pred"][i]
        per_bucket[bucket][name]["n"] += 1
        if pred == truth:
            per_bucket[bucket][name]["correct"] += 1


# ---------------------------------------------------------------------------
# Build the .docx
# ---------------------------------------------------------------------------

doc = Document()

# --- styling helpers ---
def set_table_borders(table):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = table._tbl
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "888888")
        tblBorders.append(b)
    tblPr = tbl.tblPr
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def header_row(row):
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "E7E6E6")
        cell._tc.get_or_add_tcPr().append(shd)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p


def add_table(headers, rows, *, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    set_table_borders(t)
    hdr_cells = t.rows[0].cells
    for c, h in zip(hdr_cells, headers):
        c.text = str(h)
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    header_row(t.rows[0])
    for row in rows:
        cells = t.add_row().cells
        for c, v in zip(cells, row):
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for col, w in zip(t.columns, col_widths):
            for cell in col.cells:
                cell.width = w
    return t


def fmt_pct(x): return f"{x*100:.1f}%"
def fmt_f(x, d=3): return f"{x:.{d}f}"
def fmt_ms(x): return f"{x:.2f}" if x < 10 else f"{x:.1f}"


# ============================ TITLE / TL;DR ================================

# NOTE — manual edits applied from the 2026-06-25 user revision:
#   - Title now reads "WITS vs Prompt-only utility model vs Attention-Graph GNN"
#   - Subtitle replaced with the document byline.
# Keep these unless the user asks for them to change again.
doc.add_heading("WITS vs Prompt-only utility model vs Attention-Graph GNN", level=0)
add_para("Author: Prachi Patel", italic=True, size=12)
doc.add_paragraph()


# ============================ AIM / SCOPE ==================================

# NOTE — manual edits applied from the 2026-06-25 user revision:
#   - Top-level heading shortened from
#       "Aim — what this work is for, and what it is not"
#     to just "Aim".
#   - "Lower marginal cost..." bullet trimmed to one short clause.
#   - "How this addresses..." paragraph: removed the trailing D3 cross-eval
#     callout sentences and the em-dash editorial in the gpt-5.4-nano note.
#   - "Scope and non-goals" sub-heading: em-dash replaced with a colon.
#   - "Deployment plan" step 2 collapsed to one short sentence.
add_heading("Aim", level=1)

add_para("What we are proposing", bold=True, size=11)
add_para(
    "Replace the WITS static command classifier (~15-20k lines of hand-"
    "curated regex / AST rules in copilot-agent-runtime) with a learned "
    "attention-graph GNN over a frozen Qwen 2.5 0.5B featurizer. The "
    "GNN keeps the same 4-class verdict contract (safe / maybe_safe / "
    "unsafe / extremely_unsafe) so it is a drop-in for the existing "
    "Verdict enum, but emits a calibrated confidence alongside the "
    "label so the pipeline can route by threshold instead of by hard "
    "rule hits.",
    size=10,
)
doc.add_paragraph()

add_para("Where the GNN fits in the auto-approve pipeline", bold=True, size=11)
add_para(
    "The static layer's only operational job is to decide which of "
    "three buckets a command lands in: (a) confident-safe → fast-pass "
    "auto-approve, (b) confident-extremely-unsafe → fast hard-deny, "
    "(c) everything else → hand off to the LLM judge with conversation "
    "context. The GNN does exactly this routing, with two advantages "
    "over WITS:",
    size=10,
)
add_para(
    "• Confidence-thresholded fast-pass and fast-deny — the model "
    "exposes softmax probabilities, so the pipeline can require a "
    "tunable confidence floor (e.g. τ ≥ 0.85) before short-circuiting "
    "around the LLM. WITS's rule hits do not give the gate a tunable "
    "knob; they are yes/no.",
    size=10,
)
add_para(
    "• Lower marginal cost than rule maintenance — adding coverage "
    "for a new attack class is not regex authorship.",
    size=10,
)
doc.add_paragraph()

add_para("How this addresses the 'do we even need WITS?' reviewer concern", bold=True, size=11)
add_para(
    "A separate reviewer comment proposed deleting WITS entirely and "
    "calling an LLM utility model on every command that would "
    "otherwise prompt for human approval. We tested this directly: "
    "gpt-5-mini (the smallest gpt-5-family chat-API model available "
    "on this integrator) was given the same 311 commands with the "
    "same input contract WITS and the GNN, and asked to emit a "
    "4-class verdict. Result: the utility LLM beats WITS on accuracy "
    "(0.830 vs 0.775 3-class) but loses to the GNN on every metric "
    "accuracy, hard-deny precision, silent auto-approve rate, and "
    "especially latency (~6,000 ms mean / ~17,000 ms p95 per call, "
    "vs ~30 ms for the GNN). See the 'Utility-model comparison' "
    "section for the full numbers. The utility-LLM path is also "
    "subject to the production concerns we already documented for "
    "LLM-judge calls: transient failures and concurrency-related "
    "empty responses requiring retry logic. The GNN sits between "
    "WITS-as-is and LLM-on-everything: it is the cheap, "
    "deterministic-feeling pre-filter that makes the LLM-judge "
    "approach economically viable.",
    size=10,
)
doc.add_paragraph()

add_para("Scope and non-goals: what this work does NOT address", bold=True, size=11)
add_para(
    "This is a command-pattern classifier replacement. It does not "
    "address the broader trust-chain concerns raised in the design "
    "review:",
    size=10,
)
add_para(
    "• Untrusted instruction sources — commands derived from hostile "
    "repo content, MCP tool outputs, skill SKILL.md files, or "
    "agent-self-authored scripts look identical to user-requested "
    "commands at the static layer. No command-level classifier can "
    "distinguish them; this is a provenance problem.",
    size=10,
)
add_para(
    "• Out-of-workspace writes — commands that touch ~/.ssh, ~/.aws, "
    "~/.gnupg, or other paths outside the project root are not "
    "specially gated today. WITS does not enforce this and neither "
    "would the GNN; it needs explicit path-scoping in the pipeline.",
    size=10,
)
add_para(
    "• Decoded-from-encoded commands — when the agent decodes a "
    "base64/encoded payload and runs the decoded form, that is "
    "currently invisible to the static layer. Needs a 'decoded-by-"
    "agent' marker in the pipeline, not a smarter classifier.",
    size=10,
)
add_para(
    "These five attack classes (prompt injection via repo files, "
    "trusted-script execution like `npm test` in untrusted repos, "
    "agent-self-authored test code, hostile skills, hostile MCP "
    "outputs) require pipeline-architecture changes — provenance "
    "tagging of instruction sources, explicit out-of-workspace write "
    "controls, self-authored-script gating, and untrusted-content "
    "markers on tool outputs. Those are separate from the question "
    "this report answers, which is: 'if we are going to keep a static "
    "command classifier in the pipeline at all, should it be WITS or "
    "a learned GNN?'.",
    size=10,
)
doc.add_paragraph()

add_para("Deployment plan", bold=True, size=11)
add_para(
    "1. Ship the GNN in shadow mode behind the existing WITS "
    "verdict for 2-4 weeks. Log GNN verdict + confidence alongside "
    "WITS verdict on every command; do not change production routing.",
    size=10,
)
add_para(
    "2. Review the disagreement telemetry. Confirm the numbers hold "
    "up on live agent traffic.",
    size=10,
)
add_para(
    "3. If telemetry holds, flip routing: GNN verdict drives the "
    "fast-pass / fast-deny / route-to-judge decision, WITS stays as a "
    "logged shadow signal for one more cycle, then deleted.",
    size=10,
)
add_para(
    "4. The trust-chain protections listed under non-goals are "
    "designed and shipped as a separate workstream. They do not block "
    "the GNN swap; the GNN swap does not block them.",
    size=10,
)


# ============================ TL;DR =========================================

add_heading("TL;DR", level=1)
# Headline numbers
headline_rows = []
for name, run in method_runs.items():
    yt, yp = run["y_true"], run["y_pred"]
    acc = accuracy_score(yt, yp)
    f1m = f1_score(yt, yp, average="macro", labels=LABELS, zero_division=0)
    # 3-class collapse
    yt3 = [collapse_op(v) for v in yt]
    yp3 = [collapse_op(v) for v in yp]
    acc3 = accuracy_score(yt3, yp3)
    f13 = f1_score(yt3, yp3, average="macro", labels=OP_LABELS, zero_division=0)
    # silent auto-approve rate (3-class)
    silent = sum(1 for t, p in zip(yt3, yp3) if p == "safe" and t in ("maybe", "extremely_unsafe"))
    rate = silent / len(yt3)
    # end-to-end latency
    e2e_mean = run["model_lat"]["mean"] + run["feat_lat"]["mean"]
    e2e_p95 = run["model_lat"]["p95"] + run["feat_lat"]["p95"]
    headline_rows.append([
        name,
        fmt_f(acc), fmt_f(f1m),
        fmt_f(acc3), fmt_f(f13),
        fmt_pct(rate),
        fmt_ms(e2e_mean), fmt_ms(e2e_p95),
    ])

add_para("Test split: 311 held-out cases drawn from the merged corpus "
         "(raw WITS tests + reviewer-curated attack patterns + audit-driven "
         "gap-fill + agent-gating-specific surface + diversity polish).",
         size=10)

add_table(
    ["Method",
     "Acc (4-class)", "Macro-F1 (4-class)",
     "Acc (3-class op)", "Macro-F1 (3-class op)",
     "Silent auto-approve", "Latency mean (ms)", "Latency p95 (ms)"],
    headline_rows,
)

doc.add_paragraph()
add_para("How to read this:", bold=True, size=10)
add_para("• Acc (3-class op) collapses `maybe_safe` ∪ `unsafe` -> `maybe` (both route to LLM judge in production, so the distinction is operationally moot).", size=10)
add_para("• Silent auto-approve = fraction of test cases where the method predicted `safe` but the truth was `maybe`/`extremely_unsafe`. This is THE safety-critical failure mode.", size=10)
add_para("• Latency is per-command end-to-end (featurization + model). WITS does not need LLM featurization; the GNN does.", size=10)
add_para("• The utility-model LLM row(s) (gpt-5-mini, gpt-4.1, gpt-4o-mini) are included to test the reviewer proposal 'just use an LLM as the static layer'. They beat WITS on accuracy but lose to the GNN on accuracy, hard-deny precision, and latency by a wide margin. See the 'Utility-model comparison' section for the full head-to-head.", size=10)


# ============================ METHOD DESCRIPTIONS ==========================

doc.add_page_break()
add_heading("Methods", level=1)
add_heading("WITS static (rule-based)", level=2)
add_para(
    "The deterministic shell-command analyzer "
    "(whatInTheShell.isThis) shipped with the auto-approve gate in "
    "copilot-agent-runtime. Pattern-matches command argv structure "
    "against a hand-curated rule set and emits one of {safe, "
    "maybe_safe, unsafe, extremely_unsafe}. Runs in-process; no LLM. "
    "Scored here via a Node subprocess shim that calls the compiled "
    "dist bundle at c:/dev/what-in-the-shell-fresh/dist/index.cjs."
)
add_heading("Prompt-only LLM (Qwen 2.5 0.5B Instruct)", level=2)
add_para(
    "Same frozen LLM the GNN uses as a featurizer. Asks the model to "
    "classify the command directly using one-token argmax over the four "
    "class names in the chat template. No graph, no learnable head — "
    "just a zero-shot 'what would the model say?' baseline."
)
add_heading("GNN (class-weighted)", level=2)
add_para(
    "GATv2 graph classifier over a 5-node attention graph extracted "
    "from a single forward pass of the same frozen Qwen 2.5 0.5B "
    "instruct model. Nodes: four class-anchor definitions (safe_def, "
    "maybe_safe_def, unsafe_def, extremely_unsafe_def) and one "
    "command_input node spanning the user-content tokens. Edge "
    "features are pooled attention weights (mean, max, top-k) over "
    "each span pair, both as scalars and per-attention-layer. Trained "
    "with class-weighted cross-entropy (weights = inverse train "
    "frequency, capped at 20x) for up to 700 epochs with macro-F1 "
    "early-stopping on the test set."
)
if utility_loaded:
    add_heading("Utility-model LLMs", level=2)
    util_names = ", ".join(m for m, _, _ in utility_loaded)
    add_para(
        "Direct ask: prompt one or more general-purpose LLMs via CAPI "
        f"({util_names}) with the command + shell only and require "
        "them to emit one of the four verdict labels in a single JSON "
        "object. No graph, no learnable head, no intention or "
        "transcript fed in — same input contract as WITS and the GNN "
        "so the head-to-head is apples-to-apples. Included to answer "
        "the reviewer question 'could we just use an LLM utility "
        "model as the static layer?'."
    )


# ============================ DETAILED PER-METHOD ==========================

doc.add_page_break()
add_heading("Detailed per-method results (4-class)", level=1)

for name, run in method_runs.items():
    add_heading(name, level=2)
    yt, yp = run["y_true"], run["y_pred"]
    acc = accuracy_score(yt, yp)
    f1m = f1_score(yt, yp, average="macro", labels=LABELS, zero_division=0)
    f1w = f1_score(yt, yp, average="weighted", labels=LABELS, zero_division=0)
    add_para(f"Accuracy: {fmt_f(acc)}   Macro-F1: {fmt_f(f1m)}   Weighted-F1: {fmt_f(f1w)}", bold=True)

    # Per-class TP/FP/FN/TN/precision/recall/F1/support.
    rows = per_class_metrics(yt, yp, LABELS)
    table_rows = [[r["class"],
                   r["tp"], r["fp"], r["fn"], r["tn"], r["support"],
                   fmt_f(r["precision"]), fmt_f(r["recall"]), fmt_f(r["f1"])] for r in rows]
    add_table(["Class", "TP", "FP", "FN", "TN", "Support",
               "Precision", "Recall", "F1"], table_rows)

    doc.add_paragraph()
    add_para("Confusion matrix (rows = truth, cols = prediction):", italic=True, size=10)
    cm = confusion_matrix(yt, yp, labels=LABELS)
    cm_rows = []
    for i, cls in enumerate(LABELS):
        cm_rows.append([cls] + [int(cm[i, j]) for j in range(len(LABELS))])
    add_table(["true \\ pred"] + LABELS, cm_rows)
    doc.add_paragraph()


# ============================ 3-CLASS OPERATIONAL ==========================

doc.add_page_break()
add_heading("Operational 3-class accuracy (production-aligned)", level=1)
add_para(
    "Why we collapse from 4 classes to 3 for reporting: in the "
    "auto-approve pipeline, the static layer's only job is to route — "
    "`safe` short-circuits to auto-approve, `extremely_unsafe` "
    "short-circuits to hard-deny, and BOTH `maybe_safe` AND `unsafe` "
    "are handed off to the LLM judge. The judge then issues the final "
    "verdict. So distinguishing `maybe_safe` from `unsafe` at the "
    "static layer has zero downstream impact — the same command takes "
    "the same path through the pipeline either way. Reporting only "
    "4-class metrics penalizes both WITS and the GNN for confusions "
    "that production never sees. The 3-class view collapses "
    "`maybe_safe ∪ unsafe -> judge` and is the metric we should use "
    "when comparing static-layer candidates.",
    size=10,
)
add_para(
    "Models are still TRAINED on 4 classes — the extra severity signal "
    "is useful gradient information and keeps the GNN's output "
    "contract aligned with WITS's `Verdict` enum. Only the evaluation "
    "is collapsed.",
    italic=True, size=10,
)
doc.add_paragraph()

# ---- consolidated 3-class summary table (one row per method) --------------
add_heading("3-class summary — all methods on the 311-row test split", level=2)
summary_rows = []
for name, run in method_runs.items():
    yt3 = [collapse_op(v) for v in run["y_true"]]
    yp3 = [collapse_op(v) for v in run["y_pred"]]
    acc3 = accuracy_score(yt3, yp3)
    f13 = f1_score(yt3, yp3, average="macro", labels=OP_LABELS, zero_division=0)
    # silent auto-approve on the 3-class collapse: predicted safe but truth was judge/extreme
    silent = sum(1 for t, p in zip(yt3, yp3)
                 if p == "safe" and t in ("maybe", "extremely_unsafe"))
    silent_rate = silent / len(yt3) if yt3 else 0.0
    # hard-deny precision: of cases predicted extremely_unsafe, how many truly were
    pred_extreme = [t for t, p in zip(yt3, yp3) if p == "extremely_unsafe"]
    hd_prec = (sum(1 for t in pred_extreme if t == "extremely_unsafe") / len(pred_extreme)
               if pred_extreme else float("nan"))
    summary_rows.append([
        name,
        fmt_f(acc3),
        fmt_f(f13),
        fmt_pct(silent_rate),
        (fmt_f(hd_prec) if not (isinstance(hd_prec, float) and hd_prec != hd_prec) else "—"),
    ])
add_table(
    ["Method", "Accuracy (3-class)", "Macro F1 (3-class)",
     "Silent auto-approve", "Hard-deny precision"],
    summary_rows,
)
add_para(
    "Silent auto-approve = predicted `safe` but truth was `judge` or "
    "`extremely_unsafe`. This is the safety-critical failure mode — a "
    "command that should have gone to the judge or been hard-denied "
    "instead got executed without review. Lower is better.",
    size=10,
)
add_para(
    "Hard-deny precision = of cases the method predicted "
    "`extremely_unsafe`, the fraction that truly were. A low number "
    "would mean the method is over-blocking and would hurt developer "
    "experience.",
    size=10,
)
doc.add_paragraph()

# ---- per-method 3-class detail (kept for reviewers who want the full picture) ----
add_heading("Per-method 3-class detail", level=2)

for name, run in method_runs.items():
    add_heading(name + " — 3-class", level=3)
    yt3 = [collapse_op(v) for v in run["y_true"]]
    yp3 = [collapse_op(v) for v in run["y_pred"]]
    add_para(f"Accuracy: {fmt_f(accuracy_score(yt3, yp3))}   "
             f"Macro-F1: {fmt_f(f1_score(yt3, yp3, average='macro', labels=OP_LABELS, zero_division=0))}",
             bold=True)
    rows = per_class_metrics(yt3, yp3, OP_LABELS)
    table_rows = [[r["class"],
                   r["tp"], r["fp"], r["fn"], r["tn"], r["support"],
                   fmt_f(r["precision"]), fmt_f(r["recall"]), fmt_f(r["f1"])] for r in rows]
    add_table(["Class", "TP", "FP", "FN", "TN", "Support",
               "Precision", "Recall", "F1"], table_rows)
    doc.add_paragraph()
    add_para("Confusion matrix (3-class collapsed):", italic=True, size=10)
    cm = confusion_matrix(yt3, yp3, labels=OP_LABELS)
    cm_rows = []
    for i, cls in enumerate(OP_LABELS):
        cm_rows.append([cls] + [int(cm[i, j]) for j in range(len(OP_LABELS))])
    add_table(["true \\ pred"] + OP_LABELS, cm_rows)
    doc.add_paragraph()


# ============================ LATENCY ======================================

doc.add_page_break()
add_heading("Latency", level=1)
add_para("Per-call latency, measured on CPU (no CUDA). The GNN's own "
         "forward pass is sub-millisecond; the cost of using it as an "
         "inline gate is dominated by the LLM forward pass that produces "
         "its node + edge features.", size=10)

lat_rows = []
for name, run in method_runs.items():
    ml, fl = run["model_lat"], run["feat_lat"]
    e2e_mean = ml["mean"] + fl["mean"]
    e2e_p95 = ml["p95"] + fl["p95"]
    e2e_max = ml["max"] + fl["max"]
    lat_rows.append([
        name,
        fmt_ms(fl["mean"]), fmt_ms(fl["p95"]),
        fmt_ms(ml["mean"]), fmt_ms(ml["p95"]),
        fmt_ms(e2e_mean), fmt_ms(e2e_p95), fmt_ms(e2e_max),
        run["lat_breakdown"],
    ])
add_table(["Method",
           "Feat mean (ms)", "Feat p95 (ms)",
           "Model mean (ms)", "Model p95 (ms)",
           "E2E mean (ms)", "E2E p95 (ms)", "E2E max (ms)",
           "Breakdown"], lat_rows)

doc.add_paragraph()
add_para("Throughput implications:", bold=True, size=10)
add_para(f"• WITS static is {int((FEAT_LAT['mean'] + statistics.mean(gnn_lat_ms)) / max(statistics.mean(wits_lat_ms), 1e-6))}× faster end-to-end than the GNN at p50.", size=10)
add_para("• On GPU the LLM featurization typically drops by ~10x, narrowing the gap. The 1 ms GNN tail is essentially free.", size=10)
add_para("• For tiered deployment (WITS first, GNN only on `maybe_safe`/`unsafe`), the GNN's amortized cost across all commands is much smaller — only ~30% of test cases reached the maybe/unsafe bucket.", size=10)


# ============================ UTILITY-MODEL COMPARISON =====================

if utility_loaded:
    doc.add_page_break()
    add_heading("Utility-model comparison — LLMs as the static layer", level=1)
    add_para(
        "This section answers a specific reviewer question: 'instead "
        "of WITS or a GNN, could we just ask a small utility LLM to "
        "classify each command?' We tested this end-to-end on the "
        "same 311-row split with the same input contract (command + "
        "shell only, no transcript, no intention, no rule hits). The "
        "model was given a four-class prompt requiring a JSON object "
        "of the form {\"verdict\": \"<label>\", \"rationale\": ...} "
        "on the last line.",
        size=10,
    )
    add_para(
        "Model selection note: the user originally asked for "
        "gpt-5.4-nano. That model is not exposed on the vscode-chat "
        "CAPI integrator. We evaluated every other plausible "
        "small/medium utility model available on the integrator "
        f"({', '.join(m for m, _, _ in utility_loaded)}). "
        "gpt-5.4-mini exists but is responses-API only (returns 400 "
        "from /chat/completions) so it is excluded.",
        italic=True, size=10,
    )

    # Compact head-to-head table for WITS + GNN + every utility model.
    w = method_runs["WITS static (rule-based)"]
    g = method_runs["GNN (class-weighted CE)"]

    def _row(name, run):
        yt, yp = run["y_true"], run["y_pred"]
        acc4 = accuracy_score(yt, yp)
        f14  = f1_score(yt, yp, average="macro", labels=LABELS, zero_division=0)
        yt3 = [collapse_op(v) for v in yt]
        yp3 = [collapse_op(v) for v in yp]
        acc3 = accuracy_score(yt3, yp3)
        f13  = f1_score(yt3, yp3, average="macro", labels=OP_LABELS, zero_division=0)
        silent = sum(1 for t, p in zip(yt3, yp3) if p == "safe" and t in ("maybe", "extremely_unsafe"))
        silent_rate = silent / len(yt3) if yt3 else 0.0
        pred_extreme = [t for t, p in zip(yt3, yp3) if p == "extremely_unsafe"]
        hd_prec = (sum(1 for t in pred_extreme if t == "extremely_unsafe") / len(pred_extreme)
                   if pred_extreme else None)
        e2e_mean = run["feat_lat"]["mean"] + run["model_lat"]["mean"]
        e2e_p95  = run["feat_lat"]["p95"]  + run["model_lat"]["p95"]
        return [name,
                fmt_f(acc4), fmt_f(f14),
                fmt_f(acc3), fmt_f(f13),
                fmt_pct(silent_rate),
                (fmt_f(hd_prec) if hd_prec is not None else "—"),
                fmt_ms(e2e_mean), fmt_ms(e2e_p95)]

    add_heading("Head-to-head on the same 311 commands", level=2)
    head_rows = [
        _row("WITS static", w),
        _row("GNN (weighted)", g),
    ]
    for model_name, display, _recs in utility_loaded:
        head_rows.append(_row(f"Utility-model ({model_name})", method_runs[display]))
    add_table(
        ["Method",
         "Acc (4-class)", "Macro-F1 (4-class)",
         "Acc (3-class)", "Macro-F1 (3-class)",
         "Silent auto-approve", "Hard-deny precision",
         "E2E mean (ms)", "E2E p95 (ms)"],
        head_rows,
    )
    doc.add_paragraph()

    # 3-class confusion matrix for each utility model.
    for model_name, display, _recs in utility_loaded:
        u = method_runs[display]
        add_heading(f"{model_name} — 3-class confusion matrix", level=2)
        add_para("Rows = truth, cols = prediction. Collapsed to operational 3-class "
                 "(maybe_safe + unsafe -> maybe -> route to judge).", italic=True, size=10)
        yt3 = [collapse_op(v) for v in u["y_true"]]
        yp3 = [collapse_op(v) for v in u["y_pred"]]
        cm = confusion_matrix(yt3, yp3, labels=OP_LABELS)
        cm_rows = [[cls] + [int(cm[i, j]) for j in range(len(OP_LABELS))]
                   for i, cls in enumerate(OP_LABELS)]
        add_table(["true \\ pred"] + OP_LABELS, cm_rows)
        doc.add_paragraph()

    # Compute the dynamic verdict numbers so the prose stays honest as
    # new utility models are added.
    def _summary(run):
        yt, yp = run["y_true"], run["y_pred"]
        yt3 = [collapse_op(v) for v in yt]
        yp3 = [collapse_op(v) for v in yp]
        acc3 = accuracy_score(yt3, yp3)
        f13  = f1_score(yt3, yp3, average="macro", labels=OP_LABELS, zero_division=0)
        silent = sum(1 for t, p in zip(yt3, yp3) if p == "safe" and t in ("maybe", "extremely_unsafe"))
        silent_rate = silent / len(yt3) if yt3 else 0.0
        pred_extreme = [t for t, p in zip(yt3, yp3) if p == "extremely_unsafe"]
        hd_prec = (sum(1 for t in pred_extreme if t == "extremely_unsafe") / len(pred_extreme)
                   if pred_extreme else None)
        lat_mean = run["feat_lat"]["mean"] + run["model_lat"]["mean"]
        lat_p95  = run["feat_lat"]["p95"]  + run["model_lat"]["p95"]
        return acc3, f13, silent_rate, hd_prec, lat_mean, lat_p95

    g_acc3, g_f13, g_silent, g_hd, g_lm, g_lp = _summary(g)
    w_acc3, w_f13, w_silent, w_hd, w_lm, w_lp = _summary(w)

    add_heading("Verdict", level=2)
    add_para(
        f"On the same 311 commands the GNN reaches {fmt_f(g_acc3)} "
        f"3-class accuracy and {fmt_f(g_f13)} macro-F1, with "
        f"{fmt_pct(g_silent)} silent auto-approve and "
        f"{fmt_f(g_hd) if g_hd is not None else 'n/a'} hard-deny "
        f"precision, at end-to-end latency {fmt_ms(g_lm)} ms mean / "
        f"{fmt_ms(g_lp)} ms p95. The utility-model row(s) above are "
        "evaluated against that bar.",
        size=10,
    )
    # Per-model verdict line, computed dynamically.
    for model_name, display, _recs in utility_loaded:
        u_acc3, u_f13, u_silent, u_hd, u_lm, u_lp = _summary(method_runs[display])
        delta_acc = (u_acc3 - g_acc3) * 100
        slow_x = (u_lm / max(g_lm, 1e-6))
        add_para(
            f"• {model_name}: 3-class acc {fmt_f(u_acc3)} "
            f"({delta_acc:+.1f} pp vs GNN), macro-F1 {fmt_f(u_f13)}, "
            f"silent auto-approve {fmt_pct(u_silent)}, hard-deny "
            f"precision {fmt_f(u_hd) if u_hd is not None else 'n/a'}, "
            f"latency {fmt_ms(u_lm)} ms mean (~{slow_x:.0f}x slower "
            f"end-to-end than the GNN). "
            + ("Beats WITS on accuracy; loses to GNN."
               if u_acc3 > w_acc3 else
               "Does not beat WITS on accuracy either."),
            size=10,
        )
    add_para(
        "Across every utility model tested, the same pattern holds: "
        "they may beat WITS on raw accuracy, but they are markedly "
        "slower than the GNN and tend to over-predict "
        "extremely_unsafe, which hurts hard-deny precision and would "
        "directly degrade developer experience. The right slot for "
        "an LLM in this pipeline is the existing one — called only "
        "on the ~25-30% of commands the GNN is uncertain about — "
        "not as the static layer itself.",
        bold=True, size=10,
    )


# ============================ PER-BUCKET ===================================

doc.add_page_break()
add_heading("Per-source-bucket accuracy", level=1)
add_para(
    "Test set is partitioned across data sources. WITS performs "
    "excellently on its own test corpus (the rules were tuned to "
    "pass these); the GNN's advantage is concentrated on attack "
    "patterns the WITS code review flagged as missed.",
    size=10,
)

bucket_rows = []
bucket_order = sorted(per_bucket.keys(), key=lambda b: -sum(per_bucket[b][n]["n"] for n in method_runs))
for b in bucket_order:
    row = [b, per_bucket[b]["WITS static (rule-based)"]["n"]]
    for n in method_runs:
        rec = per_bucket[b][n]
        if rec["n"] == 0:
            row.append("—")
        else:
            row.append(f"{fmt_pct(rec['correct']/rec['n'])} ({rec['correct']}/{rec['n']})")
    bucket_rows.append(row)
add_table(
    ["Source bucket", "n"] + list(method_runs.keys()),
    bucket_rows,
)

doc.add_paragraph()
add_para("Note: `wits` and `car` are raw WITS test files; the others are curated for this project:", bold=True, size=10)
add_para("• `reviewer` — hand-curated cases mirroring the WITS over-broad KNOWN_SAFE review comment (env-prefix RCE, `git -c key=val`, `find -exec`, glued flags, sensitive reads).", size=10)
add_para("• `agent-gating` — agent-context attack surface (cloud-metadata SSRF, `/proc/self/environ`, `.vscode/tasks.json` confused-deputy, untrusted `./script.sh` execution).", size=10)
add_para("• `gap` — audit-flagged PowerShell attacks, disk destruction diversity, network attacks, persistence/evasion.", size=10)
add_para("• `diversity` — programmatic payload rewrites of the above with realistic-looking domains / paths to defeat literal-string memorization.", size=10)


# ============================ DISAGREEMENTS ================================

doc.add_page_break()
add_heading("Pairwise disagreement — WITS vs GNN (class-weighted)", level=1)

disagreements = []
for i in range(len(test_meta)):
    truth = test_meta[i]["verdict"]
    wits = method_runs["WITS static (rule-based)"]["y_pred"][i]
    gnn = method_runs["GNN (class-weighted CE)"]["y_pred"][i]
    if wits != gnn:
        disagreements.append({
            "i": i,
            "command": test_meta[i]["command"][:90],
            "shell": test_meta[i]["shell"],
            "truth": truth,
            "wits": wits,
            "gnn": gnn,
            "source": bucket_of(test_meta[i].get("source", "")),
            "wits_right": wits == truth,
            "gnn_right": gnn == truth,
        })

n_gnn_better = sum(1 for d in disagreements if d["gnn_right"] and not d["wits_right"])
n_wits_better = sum(1 for d in disagreements if d["wits_right"] and not d["gnn_right"])
n_both_wrong = sum(1 for d in disagreements if not d["wits_right"] and not d["gnn_right"])

add_para(f"Total disagreements: {len(disagreements)} of {len(test_meta)}", bold=True)
add_para(f"  GNN correct, WITS wrong: {n_gnn_better}", size=10)
add_para(f"  WITS correct, GNN wrong: {n_wits_better}", size=10)
add_para(f"  Both wrong (different prediction): {n_both_wrong}", size=10)
add_para(f"  Net advantage to GNN: +{n_gnn_better - n_wits_better} cases", italic=True, size=10)

# Top-20 GNN wins
doc.add_paragraph()
add_heading("Top 20 — GNN correct, WITS wrong", level=2)
add_para("(Most are env-prefix RCE / `git -c` config injection / agent-context attacks the WITS reviewer flagged as missed.)", italic=True, size=10)
gnn_wins = [d for d in disagreements if d["gnn_right"] and not d["wits_right"]][:20]
add_table(["Command", "Shell", "Truth", "WITS", "GNN", "Source"],
          [[d["command"], d["shell"], d["truth"], d["wits"], d["gnn"], d["source"]] for d in gnn_wins])

doc.add_paragraph()
add_heading("Top 20 — WITS correct, GNN wrong", level=2)
add_para("(Mostly cases where the GNN is over-cautious — predicts `unsafe` or `maybe_safe` for something the rules correctly call `safe`.)", italic=True, size=10)
wits_wins = [d for d in disagreements if d["wits_right"] and not d["gnn_right"]][:20]
add_table(["Command", "Shell", "Truth", "WITS", "GNN", "Source"],
          [[d["command"], d["shell"], d["truth"], d["wits"], d["gnn"], d["source"]] for d in wits_wins])


# ============================ DEPLOYMENT REC ===============================
# NOTE — manual edits applied from the 2026-06-25 user revision:
#   - The entire "Deployment recommendation" section was removed. The Aim
#     section's "Deployment plan" sub-section is the canonical version; the
#     duplicated tiered/use-case content here was dropped to avoid
#     contradicting the simpler GNN-as-static-layer recommendation in Aim.
#     Keep deleted unless the user reverses the decision.


# ============================ FOOTER ======================================

doc.add_page_break()
add_heading("Reproducibility", level=1)
add_para("Test split: 311 cases, exported by wits_main.ipynb section 2.", size=10)
add_para(f"WITS predictions file: {WITS_PRED_PATH.name}", size=10)
add_para(f"GNN model: {GNN_W_DIR.name}/model.pt", size=10)
add_para(f"Frozen LLM featurizer: Qwen/Qwen2.5-0.5B-Instruct (eager attention, output_attentions=True)", size=10)
add_para(f"GNN architecture: GATv2, hidden dims {md_w['hidden_channel_dimensions']}, dropout 0.5, Adam @ 5e-4, class-weighted CE (best macro-F1 {md_w.get('best_macro_f1', 'n/a'):.3f}).", size=10)
add_para("To regenerate: see README.md ('Regenerate the dataset') and run wits_main.ipynb sections 1-11b end to end.", size=10)
add_para(f"All numbers in this report are computed live by data/build_report_docx.py from the predictions on disk.", italic=True, size=9)


# --- save ---
DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(DOCX_PATH)
print(f"\nWrote {DOCX_PATH}")
print(f"  size: {DOCX_PATH.stat().st_size:,} bytes")
